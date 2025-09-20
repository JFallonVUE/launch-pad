import os
from datetime import datetime
from typing import Dict, Any, List
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_BREAK
from app.config import settings

HEADS = [
    "Proposal + Listing Lingo Pack",
    "Chosen Services & Rationale",
    "I. Core Listing & Print",
    "II. Digital & Social",
    "III. Direct Outreach",
    "Phase I: Core Asset Creation (Pre-Launch)",
    "Phase II: Release & Distribution (Go-Live)",
    "Operational Checklists",
    "Week-1 Cadence (time windows only)",
    "KPIs",
    "Compliance"
]

def H(doc: Document, text: str, size: int=16):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)

def L(doc: Document, text: str):
    doc.add_paragraph(f"• {text}")

def build_proposal_docx(
    answers: Dict[str, Any],
    stacks: Dict[str, Any],
    bias_plans: Dict[str, Any],
    chosen_tier: str,
    chosen_bias_key: str,
    copy_pack: Dict[str, Any],
) -> str:
    os.makedirs(settings.exports_dir, exist_ok=True)
    doc = Document()
    H(doc, HEADS[0], 20)
    doc.add_paragraph(f"Mode: Deep Dive    Created: {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}")
    doc.add_paragraph(f"Chosen Tier: {chosen_tier}    Chosen Bias: {chosen_bias_key}")

    # Services section
    H(doc, HEADS[1], 16)
    chosen = next((s for s in stacks["stacks"] if s["tier"]==chosen_tier), stacks["stacks"][0])
    for s in chosen["services"]:
        name = s.get("name", s.get("service_id"))
        doc.add_paragraph(f"• {name} — {s.get('rationale','Selected for fit.')}")
    doc.add_paragraph()

    # Listing Lingo
    H(doc, HEADS[2]); pack = copy_pack.get("core_listing_print",{})
    doc.add_paragraph(pack.get("mls_description",""))
    doc.add_paragraph(f"Flyer: {pack.get('flyer_headline','')}")
    doc.add_paragraph(pack.get("flyer_short",""))
    for b in pack.get("bulleted_specs",[]):
        L(doc,b)

    H(doc, HEADS[3]); d = copy_pack.get("digital_social",{})
    for k in ["just_listed","open_house","feature_highlights","under_contract_sold","ads"]:
        v = d.get(k,[])
        if isinstance(v, list):
            for item in v: L(doc, item)
    vids = d.get("video_scripts",{})
    if vids:
        L(doc, f"Walkthrough video: {vids.get('walkthrough','')}")
        L(doc, f"Reels/Shorts: {vids.get('reels','')}")

    H(doc, HEADS[4]); out = copy_pack.get("direct_outreach",{})
    L(doc, out.get("email_blast",""))
    tpls = out.get("inquiry_templates",{})
    for k,v in tpls.items(): L(doc, f"{k.upper()}: {v}")
    L(doc, out.get("oh_followup",""))

    H(doc, HEADS[5]); p1 = copy_pack.get("phases",{}).get("phase_i_prelaunch",[])
    for i in p1: L(doc,i)

    H(doc, HEADS[6]); p2 = copy_pack.get("phases",{}).get("phase_ii_golive",[])
    for i in p2: L(doc,i)

    H(doc, HEADS[7]); chk = copy_pack.get("checklists",{})
    for name, items in chk.items():
        doc.add_paragraph(name.replace("_"," ").title())
        for it in items: L(doc, it)

    H(doc, HEADS[8]); for w in copy_pack.get("week1_cadence",[]): L(doc, w)
    H(doc, HEADS[9]); for k in copy_pack.get("kpis",[]): L(doc,k)
    H(doc, HEADS[10]); doc.add_paragraph(copy_pack.get("compliance","Keep copy neutral and factual."))

    fname = f"proposal_{chosen_tier.lower()}_{chosen_bias_key}_{int(datetime.utcnow().timestamp())}.docx"
    doc.save(os.path.join(settings.exports_dir, fname))
    return fname
