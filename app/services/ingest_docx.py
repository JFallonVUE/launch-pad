"""
Builds normalized JSON KB for services/biases and (re)indexes the vector store.

- On startup, if data/catalog.json and data/biases.json already exist, we keep them.
- If DOCX/XLSX in data/source/ exist, we (re)materialize JSON from those.
- POST /admin/reload-kb forces re-materialization if sources present.
"""
import os, json, re
from typing import Dict, Any, List
from app.config import settings

DATA_DIR = "data"
SRC_DIR = os.path.join(DATA_DIR, "source")
CATALOG_JSON = os.path.join(DATA_DIR, "catalog.json")
BIASES_JSON = os.path.join(DATA_DIR, "biases.json")

def _file_exists(name: str) -> bool:
    return os.path.exists(os.path.join(SRC_DIR, name))

def _read_prices_xlsx():
    # Lightweight XLSX parse via csv fallback if needed; prefer pandas if installed
    try:
        import pandas as pd
        p = os.path.join(SRC_DIR, "VUE_Services_SKU_Prices.xlsx")
        if os.path.exists(p):
            df = pd.read_excel(p)
            # Expect columns: Category, Service Name, SKU, SKU3, Price
            records = []
            for _, r in df.iterrows():
                records.append({
                    "category": f"{r.get('Category','')}".strip(),
                    "name": f"{r.get('Service Name','')}".strip(),
                    "sku": f"{r.get('SKU','')}".strip(),
                    "sku3": f"{r.get('SKU3','')}".strip(),
                    "price": float(str(r.get('Price','0')).replace("$","").replace(",","")) if pd.notna(r.get('Price')) else None
                })
            return records
    except Exception:
        pass
    return []

def _extract_services_from_docx() -> List[Dict[str, Any]]:
    # Best-effort parse if python-docx is available
    services = []
    doc_path = os.path.join(SRC_DIR, "VUE Services 2026.docx")
    if not os.path.exists(doc_path):
        return services
    try:
        from docx import Document
        d = Document(doc_path)
        text = "\n".join(p.text for p in d.paragraphs)
    except Exception:
        # fallback: no docx lib
        return services

    # Very light, heuristic extraction (headings and blocks)
    blocks = [b.strip() for b in re.split(r"\n{2,}", text) if b.strip()]
    for b in blocks:
        if "Price" in b or "Deliverables" in b or "The marketing logic" in b:
            # naive sectioning; we will still rely on XLSX for SKUs & prices
            title = b.splitlines()[0].strip()
            description = "\n".join(b.splitlines()[1:])
            services.append({
                "service_id": re.sub(r"[^A-Za-z0-9]+","-", title.lower()).strip("-"),
                "name": title,
                "category": "Unknown",
                "description": description,
                "deliverables": [],
                "constraints": [],
                "compatible_biases": [],
                "price_band": None
            })
    return services

def _seed_biases_from_docx() -> List[Dict[str, Any]]:
    path = os.path.join(SRC_DIR, "Biases.docx")
    items = []
    if not os.path.exists(path):
        return items
    try:
        from docx import Document
        d = Document(path)
        text = "\n".join(p.text for p in d.paragraphs)
    except Exception:
        return items
    # Heuristic: split by bias names followed by colon
    for chunk in re.split(r"\n(?=[A-Z][A-Za-z ]+:\s)", text):
        t = chunk.strip()
        if not t: continue
        parts = t.split(":", 1)
        if len(parts) != 2: continue
        name, body = parts[0].strip(), parts[1].strip()
        key = re.sub(r"[^a-z0-9]+","-", name.lower()).strip("-")
        items.append({
            "key": key,
            "name": name,
            "definition": body[:500],
            "copy_patterns": [],
            "cadence_patterns": [{"window":"Morning 9–11a"},{"window":"Lunch 12–2p"},{"window":"Evening 5–8p"}],
            "compatible_services": []
        })
    return items

def _materialize_catalog() -> Dict[str, Any]:
    """
    Merge XLSX price table with service narratives from DOCX (if available).
    For robustness, we also ship a curated baseline catalog.json in /data.
    """
    # Start with existing file if exists
    base = {}
    if os.path.exists(CATALOG_JSON):
        try:
            with open(CATALOG_JSON,"r") as f:
                base = json.load(f)
        except Exception:
            base = {}

    # augment from sources
    prices = _read_prices_xlsx()
    doc_services = _extract_services_from_docx()
    doc_by_name = {s["name"].strip().lower(): s for s in doc_services}

    services = []
    for rec in prices:
        name = rec["name"]
        from_doc = doc_by_name.get(name.lower())
        service = {
            "service_id": rec.get("sku3") or rec.get("sku"),
            "sku": rec.get("sku"),
            "sku3": rec.get("sku3"),
            "category": rec.get("category"),
            "name": name,
            "description": (from_doc["description"] if from_doc else base.get("services",{}).get(name,{}).get("description","")).strip(),
            "deliverables": base.get("services",{}).get(name,{}).get("deliverables",[]),
            "constraints": base.get("services",{}).get(name,{}).get("constraints",[]),
            "compatible_biases": base.get("services",{}).get(name,{}).get("compatible_biases",[]),
            "price": rec.get("price")
        }
        services.append(service)

    catalog = {"services": services}
    with open(CATALOG_JSON, "w") as f:
        json.dump(catalog, f, indent=2)
    return catalog

def _materialize_biases() -> Dict[str, Any]:
    base = {}
    if os.path.exists(BIASES_JSON):
        try:
            with open(BIASES_JSON,"r") as f:
                base = json.load(f)
        except Exception:
            base = {}

    doc = _seed_biases_from_docx()
    # If doc parse failed, keep base (we ship a robust biases.json already)
    result = {"biases": doc if doc else base.get("biases", [])}
    if not result["biases"]:
        result = base
    with open(BIASES_JSON, "w") as f:
        json.dump(result, f, indent=2)
    return result

def ensure_kb_materialized():
    # If both files exist, nothing to do (the shipped ones cover baseline)
    if not os.path.exists(CATALOG_JSON) or not os.path.exists(BIASES_JSON):
        os.makedirs(DATA_DIR, exist_ok=True)
        os.makedirs(SRC_DIR, exist_ok=True)
        if not os.path.exists(CATALOG_JSON):
            _materialize_catalog()
        if not os.path.exists(BIASES_JSON):
            _materialize_biases()

def rebuild_from_sources():
    os.makedirs(DATA_DIR, exist_ok=True)
    os.makedirs(SRC_DIR, exist_ok=True)
    cat = _materialize_catalog()
    biases = _materialize_biases()
    return {"catalog": len(cat.get("services",[])), "biases": len(biases.get("biases",[]))}
